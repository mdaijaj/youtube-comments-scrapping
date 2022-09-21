# pip install selenium
# pip install beautifulsoup4
# pip install webdriver-manager
from selenium import webdriver
from webdriver_manager.firefox import GeckoDriverManager
from bs4 import BeautifulSoup
import time
import docx


def ScrapComment(url):
    option = webdriver.FirefoxOptions()
    option.add_argument("--headless")
    driver = webdriver.Firefox(executable_path=GeckoDriverManager().install(), options=option)
    driver.get(url)
    pre_h = 0
    while True:
        height = driver.execute_script("""
                function getActualHeight() {
                    return Math.max(
                        Math.max(document.body.scrollHeight, document.documentElement.scrollHeight),
                        Math.max(document.body.offsetHeight, document.documentElement.offsetHeight),
                        Math.max(document.body.clientHeight, document.documentElement.clientHeight)
                    );
                }
                return getActualHeight();
            """)
        driver.execute_script(f"window.scrollTo({pre_h},{pre_h + 200})")

        # time sleep your network connection depends
        time.sleep(1)
        pre_h +=200  
        if pre_h >= height:
            break
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    driver.quit()
    title_text_div = soup.select_one('#container h1')
    title = title_text_div and title_text_div.text
    comment_div = soup.select("#content #content-text")
    comment_list = [x.text for x in comment_div]
    print(title, comment_list)


    # export docs here.
    doc= docx.Document()
    doc.add_heading('Youtube Scrapping', 1)
    doc.add_paragraph('this is some document store file')

    doc.add_heading(title, 1)
    doc.add_paragraph(comment_list, style='List Number')
    doc.save("newfile.docx")


if __name__ == "__main__":

    urls = [
        "https://www.youtube.com/watch?v=cgNQgcUgq0U",
        "https://www.youtube.com/watch?v=MkE_EwO76b0",
        'https://www.youtube.com/watch?v=XVv6mJpFOb0',
    ]

ScrapComment(urls[1])